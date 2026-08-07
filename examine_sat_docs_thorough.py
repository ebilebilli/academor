# -*- coding: utf-8 -*-
"""Thorough SAT Word document analysis (ordered body walk + answer keys)."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

MATH_DOC = Path(r"c:\Users\user\Desktop\SAT Mock\One\SAT_Math_One.docx")
MATH_KEY = Path(r"c:\Users\user\Desktop\SAT Mock\One\SAT_Math_One_Answer_Keys.docx")
VERBAL_DOC = Path(r"c:\Users\user\Desktop\SAT Mock\One\SAT_Verbal_One.docx")
VERBAL_KEY = Path(r"c:\Users\user\Desktop\SAT Mock\One\Sat_Verbal_One_Answer_Keys.docx")
OUT_PATH = Path(r"c:\Users\user\Desktop\Academor\sat_analysis_report.json")

MATH_Q_RE = re.compile(r"^(\d+)\)\s*")
VERBAL_Q_RE = re.compile(r"^Question\s+(\d+)\.?\s*$", re.IGNORECASE)
OPTION_LETTER_ONLY_RE = re.compile(r"^([A-D])[\)\.]?\s*$", re.IGNORECASE)
OPTION_WITH_TEXT_RE = re.compile(r"^([A-D])[\)\.:]\s+(.+)$", re.IGNORECASE)
MCQ_ANSWER_RE = re.compile(r"^[A-D]$", re.IGNORECASE)
MODULE_RE = re.compile(r"module\s*(\d+)", re.IGNORECASE)
CORRUPT_FRAC_RE = re.compile(r'^(\d+)"\s*>\s*$')

NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}


def qn_v(tag: str) -> str:
    prefix, local = tag.split(":")
    return "{%s}%s" % (NSMAP[prefix], local)


def iter_block_items(parent) -> Iterator[Paragraph | Table]:
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def paragraph_text(p: Paragraph) -> str:
    return (p.text or "").strip()


def blip_rids_in_element(element) -> list[str]:
    rids: list[str] = []
    for blip in element.findall(".//" + qn_v("a:blip")):
        rid = blip.get(qn_v("r:embed")) or blip.get(qn_v("r:link"))
        if rid:
            rids.append(rid)
    for im in element.findall(".//" + qn_v("v:imagedata")):
        rid = im.get(qn_v("r:id"))
        if rid:
            rids.append(rid)
    return rids


def resolve_media(doc: Document, rid: str) -> str | None:
    try:
        rel = doc.part.rels[rid]
    except KeyError:
        return None
    target = rel.target_ref
    return Path(target).name if target else None


def extract_images_from_element(doc: Document, element) -> list[str]:
    names = []
    for rid in blip_rids_in_element(element):
        name = resolve_media(doc, rid)
        if name:
            names.append(name)
    return names


def count_omml(element) -> int:
    return len(element.findall(".//" + qn_v("m:oMath")))


def serialize_omml(node) -> str:
    local = node.tag.split("}")[-1]
    if local == "t":
        return node.text or ""
    if local == "f":
        num = node.find(qn_v("m:num"))
        den = node.find(qn_v("m:den"))

        def ser_side(side):
            if side is None:
                return ""
            return "".join(serialize_omml(c) for c in side)

        return f"({ser_side(num)})/({ser_side(den)})"
    if local == "rad":
        e = node.find(qn_v("m:e"))
        return f"sqrt({''.join(serialize_omml(c) for c in e) if e is not None else ''})"
    if local == "sSup":
        e = node.find(qn_v("m:e"))
        sup = node.find(qn_v("m:sup"))
        b = "".join(serialize_omml(c) for c in e) if e is not None else ""
        s = "".join(serialize_omml(c) for c in sup) if sup is not None else ""
        return f"{b}^{s}"
    if local in (
        "rPr",
        "ctrlPr",
        "jc",
        "fPr",
        "numPr",
        "denPr",
        "radPr",
        "degHide",
        "deg",
        "sSupPr",
        "sSubPr",
        "argPr",
        "br",
    ):
        return ""
    return "".join(serialize_omml(c) for c in node)


def omml_texts_in_element(element) -> list[str]:
    out = []
    for om in element.findall(".//" + qn_v("m:oMath")):
        t = serialize_omml(om).strip()
        if t:
            out.append(t)
    return out


def recover_corrupt_fraction(plain: str) -> tuple[str, str]:
    """
    Recover mangled MathML leftovers like '3618\" >' / '310\" >' / '1517\" >'.
    Returns (answer, note).
    """
    m = CORRUPT_FRAC_RE.match(plain.strip())
    if not m:
        return plain, ""
    digits = m.group(1)
    # Prefer balanced / common SAT splits
    guesses = []
    n = len(digits)
    for i in range(1, n):
        a, b = digits[:i], digits[i:]
        if b.lstrip("0") != b and b != "0":
            # leading zeros in den unlikely except 0.x handled elsewhere
            if b.startswith("0") and len(b) > 1:
                continue
        guesses.append(f"{a}/{b}")
    # Prefer splits where both parts are reasonable lengths (1-3 digits)
    preferred = [g for g in guesses if all(1 <= len(p) <= 3 for p in g.split("/"))]
    # Known-looking: 3/10, 15/17, 36/18
    best = preferred[0] if preferred else (guesses[0] if guesses else digits)
    # Prefer specific known patterns
    for cand in ("3/10", "15/17", "36/18", "5/17", "16/18"):
        if cand in preferred or cand in guesses:
            best = cand
            break
    # digit-length heuristics
    if digits == "310":
        best = "3/10"
    elif digits == "1517":
        best = "15/17"
    elif digits == "3618":
        best = "36/18"
    return best, f"recovered_from_corrupt_mathml plain={plain!r} digits={digits}"


def extract_cell_answer(doc: Document, cell) -> tuple[str, str, str]:
    """Return (answer, kind, note)."""
    plain = cell.text.strip()
    omml_list = omml_texts_in_element(cell._tc)
    omml_joined = "".join(omml_list).strip()
    imgs = extract_images_from_element(doc, cell._tc)

    # Corrupted fraction pattern in plain text
    if CORRUPT_FRAC_RE.match(plain):
        ans, note = recover_corrupt_fraction(plain)
        return ans, "spr", note

    # Clean OMML
    if omml_joined and "mathxmlns" not in omml_joined.lower() and not omml_joined.startswith("<"):
        if MCQ_ANSWER_RE.match(omml_joined) and MCQ_ANSWER_RE.match(plain or omml_joined):
            return omml_joined.upper(), "mcq", "from_omml_letter"
        # multi-value like 15or-5
        return omml_joined, "spr", "from_omml"

    if plain:
        if MCQ_ANSWER_RE.match(plain):
            return plain.upper(), "mcq", "from_plain"
        if "mathxmlns" in plain.lower() or CORRUPT_FRAC_RE.match(plain):
            ans, note = recover_corrupt_fraction(plain)
            return ans, "spr", note or "corrupt_plain"
        return plain, "spr", "from_plain"

    if omml_joined:
        # corrupted omml only
        if "mathxmlns" in omml_joined.lower():
            return "[CORRUPT_OMML]", "spr", f"omml={omml_joined!r}"
        return omml_joined, "spr", "from_omml"

    if imgs:
        return f"[image:{','.join(imgs)}]", "spr", "from_image"
    return "", "spr", "empty"


def parse_math_answer_key(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    rows_out = []
    answer_map_by_key_id: dict[int, str] = {}
    sequential: list[dict[str, Any]] = []

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 3:
                continue
            key_id_s = cells[0].text.strip()
            subject = cells[1].text.strip()
            if not key_id_s.isdigit():
                continue
            key_id = int(key_id_s)
            ans, kind, note = extract_cell_answer(doc, cells[2])
            entry = {
                "key_id": key_id,
                "subject": subject,
                "answer": ans,
                "kind": kind,
                "note": note,
                "raw_cell_text": cells[2].text.strip(),
            }
            rows_out.append(entry)
            answer_map_by_key_id[key_id] = ans
            sequential.append(entry)

    return {
        "path": str(path),
        "rows": rows_out,
        "answer_map_by_key_id": {str(k): v for k, v in sorted(answer_map_by_key_id.items())},
        "sequential_answers": [e["answer"] for e in sequential],
        "sequential_kinds": [e["kind"] for e in sequential],
        "count": len(sequential),
        "mcq_in_key": sum(1 for e in sequential if e["kind"] == "mcq"),
        "spr_in_key": sum(1 for e in sequential if e["kind"] == "spr"),
    }


def parse_verbal_answer_key(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    module = None
    by_module: dict[str, list[str]] = {}
    sequential: list[dict[str, Any]] = []

    for block in iter_block_items(doc):
        if not isinstance(block, Paragraph):
            continue
        text = paragraph_text(block)
        if not text:
            continue
        mm = MODULE_RE.search(text)
        if mm and len(text) < 40:
            module = int(mm.group(1))
            by_module.setdefault(str(module), [])
            continue
        if MCQ_ANSWER_RE.match(text):
            ans = text.upper()
            if module is None:
                module = 1
                by_module.setdefault("1", [])
            by_module[str(module)].append(ans)
            sequential.append(
                {
                    "module": module,
                    "local_number": len(by_module[str(module)]),
                    "answer": ans,
                    "kind": "mcq",
                }
            )

    answer_map = {
        f"M{e['module']}-Q{e['local_number']}": e["answer"] for e in sequential
    }
    return {
        "path": str(path),
        "by_module": by_module,
        "sequential": sequential,
        "answer_map": answer_map,
        "count": len(sequential),
        "mcq_in_key": sum(1 for e in sequential if e["kind"] == "mcq"),
        "spr_in_key": sum(1 for e in sequential if e["kind"] == "spr"),
    }


def analyze_question_doc(
    path: Path,
    *,
    style: str,
    answers_by_module_local: dict[tuple[int, int], str] | None = None,
) -> dict[str, Any]:
    doc = Document(str(path))
    questions: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    module = 1
    last_option_letter: str | None = None
    seq_index = 0

    def start_question(local_n: int, header_text: str):
        nonlocal current, last_option_letter, seq_index
        seq_index += 1
        current = {
            "global_index": seq_index,
            "module": module,
            "number": local_n,
            "id": f"M{module}-Q{local_n}",
            "header": header_text,
            "stem_parts": [],
            "options": {},
            "option_order": [],
            "has_abcd_in_text": False,
            "images": [],
            "images_between_options": [],
            "omml_snippets": [],
            "omml_between_options": [],
            "raw_lines": [header_text],
        }
        questions.append(current)
        last_option_letter = None

    def attach_rich(element):
        nonlocal last_option_letter
        if current is None:
            return
        imgs = extract_images_from_element(doc, element)
        ommls = omml_texts_in_element(element)
        for im in imgs:
            entry = {
                "filename": im,
                "context": "after_option" if last_option_letter else "stem_or_body",
                "after_option": last_option_letter,
            }
            current["images"].append(entry)
            if last_option_letter:
                current["images_between_options"].append(
                    {"filename": im, "after_option": last_option_letter}
                )
        for om in ommls:
            current["omml_snippets"].append(
                {"text": om[:120], "after_option": last_option_letter}
            )
            if last_option_letter:
                current["omml_between_options"].append(
                    {"text": om[:80], "after_option": last_option_letter}
                )

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = paragraph_text(block)
            el = block._element

            if text:
                mm = MODULE_RE.search(text)
                if mm and (
                    "section" in text.lower()
                    or text.lower().startswith("module")
                    or "math" in text.lower()
                    or len(text) < 60
                ):
                    module = int(mm.group(1))
                    current = None
                    last_option_letter = None
                    continue

                if style == "math":
                    qm = MATH_Q_RE.match(text)
                    if qm:
                        start_question(int(qm.group(1)), text.split("\n", 1)[0][:40])
                        rest = text[qm.end() :].strip()
                        if rest:
                            current["stem_parts"].append(rest)
                            current["raw_lines"].append(text)
                        attach_rich(el)
                        continue
                else:
                    qm = VERBAL_Q_RE.match(text)
                    if qm:
                        start_question(int(qm.group(1)), text)
                        attach_rich(el)
                        continue

            if current is None:
                continue

            if text:
                current["raw_lines"].append(text)

            ol = OPTION_LETTER_ONLY_RE.match(text) if text else None
            ot = OPTION_WITH_TEXT_RE.match(text) if text else None

            if ol:
                letter = ol.group(1).upper()
                current["options"].setdefault(letter, "")
                if letter not in current["option_order"]:
                    current["option_order"].append(letter)
                current["has_abcd_in_text"] = True
                last_option_letter = letter
                attach_rich(el)
                continue

            if ot:
                letter = ot.group(1).upper()
                current["options"][letter] = ot.group(2).strip()
                if letter not in current["option_order"]:
                    current["option_order"].append(letter)
                current["has_abcd_in_text"] = True
                last_option_letter = letter
                attach_rich(el)
                continue

            if text:
                if (
                    last_option_letter is not None
                    and last_option_letter in current["options"]
                    and not current["options"][last_option_letter]
                ):
                    current["options"][last_option_letter] = text
                elif last_option_letter and len(current["option_order"]) < 4:
                    prev = current["options"].get(last_option_letter, "")
                    if prev and len(prev) < 300:
                        current["options"][last_option_letter] = (prev + " " + text).strip()
                    elif not prev:
                        current["options"][last_option_letter] = text
                    else:
                        current["stem_parts"].append(text)
                else:
                    current["stem_parts"].append(text)

            attach_rich(el)

        elif isinstance(block, Table):
            if current is None:
                continue
            for row in block.rows:
                for cell in row.cells:
                    ct = " ".join(paragraph_text(pp) for pp in cell.paragraphs).strip()
                    if ct and re.fullmatch(r"[A-D]", ct, re.I):
                        letter = ct.upper()
                        current["options"].setdefault(letter, "")
                        if letter not in current["option_order"]:
                            current["option_order"].append(letter)
                        current["has_abcd_in_text"] = True
                        last_option_letter = letter
                    elif ct:
                        ot2 = OPTION_WITH_TEXT_RE.match(ct) or OPTION_LETTER_ONLY_RE.match(ct)
                        if ot2:
                            letter = ot2.group(1).upper()
                            opt_txt = ot2.group(2).strip() if ot2.lastindex and ot2.lastindex >= 2 else ""
                            current["options"][letter] = opt_txt
                            if letter not in current["option_order"]:
                                current["option_order"].append(letter)
                            current["has_abcd_in_text"] = True
                            last_option_letter = letter
                    attach_rich(cell._tc)

    media_files = []
    with zipfile.ZipFile(path, "r") as zf:
        media_files = sorted(
            n for n in zf.namelist() if "/media/" in n.replace("\\", "/").lower()
        )

    used_images = set()
    report_questions = []
    images_per_question = []
    mcq_count = 0
    spr_count = 0
    answer_map_out: dict[str, str] = {}

    for q in questions:
        stem = " ".join(q["stem_parts"]).strip()
        stem120 = stem[:120]
        opts = q["options"]
        has_opts = bool(q["option_order"])
        opt_short = {}
        for let in q["option_order"]:
            otxt = opts.get(let, "")
            # If option text empty but OMML followed this letter, note that
            omml_for = [
                o["text"]
                for o in q["omml_between_options"]
                if o.get("after_option") == let
            ]
            img_for = [
                o["filename"]
                for o in q["images_between_options"]
                if o.get("after_option") == let
            ]
            display = otxt
            if not display and omml_for:
                display = f"[omml:{omml_for[0][:60]}]"
            if not display and img_for:
                display = f"[image:{img_for[0]}]"
            opt_short[let] = display if len(display) <= 80 else display[:77] + "..."

        ans = None
        if answers_by_module_local:
            ans = answers_by_module_local.get((q["module"], q["number"]))

        if ans is not None:
            answer_map_out[q["id"]] = ans
            ans_s = str(ans).strip()
            if MCQ_ANSWER_RE.match(ans_s):
                suggested = "mcq"
                note = (
                    "mcq_text_options"
                    if has_opts and any(opts.get(l) for l in q["option_order"])
                    else (
                        "mcq_omml_or_image_options"
                        if has_opts
                        else "mcq_answer_letter_no_options_detected"
                    )
                )
            else:
                suggested = "spr"
                note = "spr_non_letter_answer"
        else:
            if has_opts:
                suggested = "mcq"
                note = "inferred_mcq_from_options_no_key"
            else:
                suggested = "spr"
                note = "inferred_spr_no_options_no_key"

        if suggested == "mcq":
            mcq_count += 1
        else:
            spr_count += 1

        img_files = [im["filename"] for im in q["images"]]
        used_images.update(img_files)
        images_per_question.append(
            {
                "id": q["id"],
                "module": q["module"],
                "number": q["number"],
                "global_index": q["global_index"],
                "image_count": len(img_files),
                "filenames": img_files,
                "omml_count": len(q["omml_snippets"]),
            }
        )

        report_questions.append(
            {
                "id": q["id"],
                "module": q["module"],
                "number": q["number"],
                "global_index": q["global_index"],
                "stem_preview": stem120,
                "stem_full_length": len(stem),
                "has_abcd_options_in_text": has_opts,
                "options": opt_short if has_opts else {},
                "option_letters_found": q["option_order"],
                "image_filenames": img_files,
                "image_details": q["images"],
                "images_between_options": q["images_between_options"],
                "omml_count": len(q["omml_snippets"]),
                "omml_snippets": q["omml_snippets"][:12],
                "omml_between_options": q["omml_between_options"],
                "answer_from_key": ans,
                "suggested_type": suggested,
                "suggested_note": note,
            }
        )

    orphan_media = [
        Path(m).name for m in media_files if Path(m).name not in used_images
    ]

    return {
        "path": str(path),
        "media_in_package": media_files,
        "media_count": len(media_files),
        "orphan_media": orphan_media,
        "question_count": len(questions),
        "mcq_count": mcq_count,
        "spr_count": spr_count,
        "total_image_occurrences": sum(len(q["images"]) for q in questions),
        "total_omml_occurrences": sum(len(q["omml_snippets"]) for q in questions),
        "images_per_question": images_per_question,
        "answer_map": answer_map_out,
        "questions": report_questions,
    }


def main():
    print("Parsing answer keys...")
    math_key = parse_math_answer_key(MATH_KEY)
    verbal_key = parse_verbal_answer_key(VERBAL_KEY)

    print(
        f"Math answers parsed: {math_key['count']} "
        f"(MCQ={math_key['mcq_in_key']}, SPR={math_key['spr_in_key']})"
    )
    print(f"Verbal answers parsed: {verbal_key['count']}")

    math_mod_local: dict[tuple[int, int], str] = {}
    for i, ans in enumerate(math_key["sequential_answers"]):
        if i < 22:
            math_mod_local[(1, i + 1)] = ans
        else:
            math_mod_local[(2, i - 21)] = ans  # i=22 -> M2-Q1

    verbal_mod_local: dict[tuple[int, int], str] = {}
    for e in verbal_key["sequential"]:
        verbal_mod_local[(e["module"], e["local_number"])] = e["answer"]

    print("Analyzing Math document...")
    math_doc = analyze_question_doc(
        MATH_DOC, style="math", answers_by_module_local=math_mod_local
    )
    print("Analyzing Verbal document...")
    verbal_doc = analyze_question_doc(
        VERBAL_DOC, style="verbal", answers_by_module_local=verbal_mod_local
    )

    report = {
        "meta": {
            "math_doc": str(MATH_DOC),
            "math_key": str(MATH_KEY),
            "verbal_doc": str(VERBAL_DOC),
            "verbal_key": str(VERBAL_KEY),
            "notes": [
                "Math questions use 'N)' headers; Verbal uses 'Question N.'",
                "Both sections have 2 modules with renumbered questions.",
                "Math answer key table uses external ids 56-98 (43 rows) vs 44 questions; M2-Q22 has no key row.",
                "Math formulas are mostly OMML (Office Math), not images; only 8 drawings in Math doc.",
                "Three SPR key cells contain corrupted MathML leftovers; recovered as likely fractions 36/18, 3/10, 15/17.",
            ],
        },
        "math": {
            "summary": {
                "question_count": math_doc["question_count"],
                "image_occurrences": math_doc["total_image_occurrences"],
                "omml_occurrences": math_doc["total_omml_occurrences"],
                "media_files_in_package": math_doc["media_count"],
                "orphan_media": math_doc["orphan_media"],
                "mcq_count": math_doc["mcq_count"],
                "spr_count": math_doc["spr_count"],
                "answer_key_count": math_key["count"],
                "answer_key_mcq": math_key["mcq_in_key"],
                "answer_key_spr": math_key["spr_in_key"],
            },
            "answer_map": math_doc["answer_map"],
            "answer_map_by_key_id": math_key["answer_map_by_key_id"],
            "answer_key_rows": math_key["rows"],
            "images_per_question": math_doc["images_per_question"],
            "questions": math_doc["questions"],
            "document_path": math_doc["path"],
            "media_in_package": math_doc["media_in_package"],
        },
        "verbal": {
            "summary": {
                "question_count": verbal_doc["question_count"],
                "image_occurrences": verbal_doc["total_image_occurrences"],
                "omml_occurrences": verbal_doc["total_omml_occurrences"],
                "media_files_in_package": verbal_doc["media_count"],
                "orphan_media": verbal_doc["orphan_media"],
                "mcq_count": verbal_doc["mcq_count"],
                "spr_count": verbal_doc["spr_count"],
                "answer_key_count": verbal_key["count"],
            },
            "answer_map": verbal_doc["answer_map"],
            "answer_key_by_module": verbal_key["by_module"],
            "images_per_question": verbal_doc["images_per_question"],
            "questions": verbal_doc["questions"],
            "document_path": verbal_doc["path"],
            "media_in_package": verbal_doc["media_in_package"],
        },
    }

    OUT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print("\n=== SAT ANALYSIS SUMMARY ===")
    print(f"Report: {OUT_PATH}")
    ms = report["math"]["summary"]
    vs = report["verbal"]["summary"]
    print(
        f"Math: {ms['question_count']} q, {ms['image_occurrences']} images, "
        f"{ms['omml_occurrences']} OMML, MCQ={ms['mcq_count']}, SPR={ms['spr_count']}"
    )
    print(
        f"Verbal: {vs['question_count']} q, {vs['image_occurrences']} images, "
        f"MCQ={vs['mcq_count']}, SPR={vs['spr_count']}"
    )
    print("\nMath images_per_question:")
    for ipq in math_doc["images_per_question"]:
        print(
            f"  {ipq['id']}: images={ipq['image_count']}{ipq['filenames']}, omml={ipq['omml_count']}"
        )
    print("\nMath SPR answers:")
    for qid, ans in math_doc["answer_map"].items():
        if not MCQ_ANSWER_RE.match(str(ans)):
            print(f"  {qid}: {ans}")
    print("\nMissing math keys:", [q["id"] for q in math_doc["questions"] if q["id"] not in math_doc["answer_map"]])


if __name__ == "__main__":
    main()
