"""Examine SAT Word documents to understand their structure."""
import sys
from pathlib import Path
from docx import Document

# Set UTF-8 encoding for output
sys.stdout.reconfigure(encoding='utf-8')

def examine_docx(docx_path, output_file):
    """Print all text from a Word document."""
    doc = Document(docx_path)
    output_file.write(f"\n=== {docx_path.name} ===\n")
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            output_file.write(f"{i:3d}: {text}\n")
    
    # Also check tables
    if doc.tables:
        output_file.write(f"\n--- Tables in {docx_path.name} ---\n")
        for table_idx, table in enumerate(doc.tables):
            output_file.write(f"\nTable {table_idx}:\n")
            for row_idx, row in enumerate(table.rows):
                row_text = [cell.text.strip() for cell in row.cells]
                output_file.write(f"  Row {row_idx}: {row_text}\n")

def main():
    base_path = Path(r'c:\Users\user\Desktop\SAT Mock\One')
    output_path = Path('c:\\Users\\user\\Desktop\\Academor\\sat_docs_examination.txt')
    
    with open(output_path, 'w', encoding='utf-8') as f:
        # Examine Math document
        math_doc = base_path / 'SAT_Math_One.docx'
        examine_docx(math_doc, f)
        
        # Examine Math answers
        math_answers = base_path / 'SAT_Math_One_Answer_Keys.docx'
        examine_docx(math_answers, f)
        
        # Examine Verbal document
        verbal_doc = base_path / 'SAT_Verbal_One.docx'
        examine_docx(verbal_doc, f)
        
        # Examine Verbal answers
        verbal_answers = base_path / 'Sat_Verbal_One_Answer_Keys.docx'
        examine_docx(verbal_answers, f)
    
    print(f"Examination saved to {output_path}")

if __name__ == '__main__':
    main()
