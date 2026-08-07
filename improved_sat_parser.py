"""Improved parser for SAT Word documents with proper structure handling."""
import json
import re
from pathlib import Path
from docx import Document

def parse_answer_key_table(doc):
    """Parse answer key from table format."""
    answer_map = {}
    if not doc.tables:
        return answer_map
    
    table = doc.tables[0]
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 3:
            try:
                q_num = int(cells[0])
                answer = cells[2].strip().upper()
                if answer and answer in 'ABCD':
                    answer_map[q_num] = answer
            except (ValueError, IndexError):
                continue
    return answer_map

def parse_math_questions(doc, answer_map):
    """Parse SAT Math questions with numbered format (1), 2), etc)."""
    questions = []
    current_question = None
    question_number = 0
    pending_option_letter = None
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        # Detect question start (e.g., "1)", "2)")
        question_match = re.match(r'^(\d+)\)\s*(.*)', text)
        if question_match:
            # Save previous question if exists
            if current_question and current_question['question']:
                # Extract embedded options from question text
                extract_embedded_options(current_question)
                questions.append(current_question)
            
            # Start new question
            question_number = int(question_match.group(1))
            current_question = {
                'id': question_number,
                'question': question_match.group(2),
                'options': [],
                'correct': None
            }
            pending_option_letter = None
        elif current_question:
            # Check if this line is just an option letter (A, B, C, D on separate lines)
            option_match = re.match(r'^([A-D])\s*$', text, re.IGNORECASE)
            if option_match:
                pending_option_letter = option_match.group(1).upper()
            elif pending_option_letter:
                # This is the option text for the pending letter
                current_question['options'].append(text)
                pending_option_letter = None
            elif re.match(r'^[A-D]\.', text):
                # Option with text after letter (alternative format)
                option_match = re.match(r'^([A-D])\.\s*(.*)', text, re.IGNORECASE)
                if option_match:
                    option_letter = option_match.group(1).upper()
                    option_text = option_match.group(2).strip()
                    current_question['options'].append(option_text)
            else:
                # Append to question text
                if current_question['question']:
                    current_question['question'] += ' ' + text
                else:
                    current_question['question'] = text
    
    # Add last question
    if current_question and current_question['question']:
        extract_embedded_options(current_question)
        questions.append(current_question)
    
    # Map answers from answer key
    for q in questions:
        if q['id'] in answer_map:
            answer_letter = answer_map[q['id']]
            answer_index = ord(answer_letter) - ord('A')
            if 0 <= answer_index < len(q['options']):
                q['correct'] = q['options'][answer_index]
    
    return questions

def extract_embedded_options(question):
    """Extract options embedded in question text (e.g., 'A option1 B option2 C option3 D option4')."""
    text = question['question']
    
    # Pattern to match options like "A\noption1 B\noption2 C\noption3 D\noption4" or "A option1 B option2 C option3 D option4"
    # First try newline pattern
    pattern = r'([A-D])\s*\n\s*([^\n]+?)(?=\s*[A-D]\s*\n|$)'
    matches = re.findall(pattern, text, re.MULTILINE)
    
    if not matches or len(matches) < 2:
        # Try space-separated pattern
        pattern = r'([A-D])\s+([A-Z]?[^A-D]+?)(?=\s+[A-D]\s+|$)'
        matches = re.findall(pattern, text)
    
    if matches and len(matches) >= 2:
        # Extract options and clean the question text
        options = []
        clean_question = text
        for letter, option_text in matches:
            option = option_text.strip()
            if option and len(option) > 0:
                options.append(option)
            # Remove this option from the question text
            clean_question = clean_question.replace(f'{letter} {option_text}', '', 1)
            clean_question = clean_question.replace(f'{letter}\n{option_text}', '', 1)
        
        if len(options) >= 2:
            question['options'] = options
            question['question'] = clean_question.strip()

def parse_verbal_questions(doc, answer_map):
    """Parse SAT Verbal questions with 'Question N.' format."""
    questions = []
    current_question = None
    question_number = 0
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        # Detect question start (e.g., "Question 1.")
        question_match = re.match(r'^Question\s+(\d+)\.\s*(.*)', text, re.IGNORECASE)
        if question_match:
            # Save previous question if exists
            if current_question and current_question['question']:
                questions.append(current_question)
            
            # Start new question
            question_number = int(question_match.group(1))
            current_question = {
                'id': question_number,
                'question': question_match.group(2),
                'options': [],
                'correct': None
            }
        elif current_question:
            # Check if this line is an answer option
            option_match = re.match(r'^([A-D])\.\s*(.*)', text, re.IGNORECASE)
            if option_match:
                option_letter = option_match.group(1).upper()
                option_text = option_match.group(2).strip()
                current_question['options'].append(option_text)
            else:
                # Append to question text
                if current_question['question']:
                    current_question['question'] += ' ' + text
                else:
                    current_question['question'] = text
    
    # Add last question
    if current_question and current_question['question']:
        questions.append(current_question)
    
    # Map answers from answer key
    for q in questions:
        if q['id'] in answer_map:
            answer_letter = answer_map[q['id']]
            answer_index = ord(answer_letter) - ord('A')
            if 0 <= answer_index < len(q['options']):
                q['correct'] = q['options'][answer_index]
    
    return questions

def main():
    base_path = Path(r'c:\Users\user\Desktop\SAT Mock\One')
    
    # Process SAT Math One
    print("Processing SAT Math One...")
    math_doc = Document(base_path / 'SAT_Math_One.docx')
    math_answers_doc = Document(base_path / 'SAT_Math_One_Answer_Keys.docx')
    
    math_answer_map = parse_answer_key_table(math_answers_doc)
    print(f"Found {len(math_answer_map)} answer keys for Math")
    
    math_questions = parse_math_questions(math_doc, math_answer_map)
    print(f"Parsed {len(math_questions)} Math questions")
    
    # Determine SAT section for Math (could be algebra or geometry_data)
    # For now, using algebra as default
    math_json = {
        'title': 'SAT Math One',
        'category_name': 'SAT Math',
        'service': 'sat',
        'is_sat': True,
        'sat_section': 'algebra',
        'questions': math_questions
    }
    
    output_path = Path('c:\\Users\\user\\Desktop\\Academor\\academor\\portals\\resources\\quiz_questions\\sat_math_one.json')
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(math_json, f, indent=2, ensure_ascii=False)
    print(f"Saved Math questions to {output_path}")
    
    # Process SAT Verbal One
    print("\nProcessing SAT Verbal One...")
    verbal_doc = Document(base_path / 'SAT_Verbal_One.docx')
    verbal_answers_doc = Document(base_path / 'Sat_Verbal_One_Answer_Keys.docx')
    
    verbal_answer_map = parse_answer_key_table(verbal_answers_doc)
    print(f"Found {len(verbal_answer_map)} answer keys for Verbal")
    
    verbal_questions = parse_verbal_questions(verbal_doc, verbal_answer_map)
    print(f"Parsed {len(verbal_questions)} Verbal questions")
    
    # Determine SAT section for Verbal (could be reading or writing)
    # For now, using reading as default
    verbal_json = {
        'title': 'SAT Verbal One',
        'category_name': 'SAT Verbal',
        'service': 'sat',
        'is_sat': True,
        'sat_section': 'reading',
        'questions': verbal_questions
    }
    
    output_path = Path('c:\\Users\\user\\Desktop\\Academor\\academor\\portals\\resources\\quiz_questions\\sat_verbal_one.json')
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(verbal_json, f, indent=2, ensure_ascii=False)
    print(f"Saved Verbal questions to {output_path}")
    
    # Print summary
    print("\n=== SUMMARY ===")
    print(f"Math: {len(math_questions)} questions")
    print(f"Verbal: {len(verbal_questions)} questions")
    
    # Note questions that might need images (those with references to figures, graphs, etc.)
    print("\n=== QUESTIONS THAT MAY NEED IMAGES ===")
    for q in math_questions:
        if any(keyword in q['question'].lower() for keyword in ['figure', 'graph', 'diagram', 'chart', 'line graph', 'dot plot']):
            try:
                print(f"Math Q{q['id']}: {q['question'][:80]}...")
            except UnicodeEncodeError:
                print(f"Math Q{q['id']}: [Contains special characters]")
    
    for q in verbal_questions:
        if any(keyword in q['question'].lower() for keyword in ['figure', 'graph', 'diagram', 'chart']):
            try:
                print(f"Verbal Q{q['id']}: {q['question'][:80]}...")
            except UnicodeEncodeError:
                print(f"Verbal Q{q['id']}: [Contains special characters]")

if __name__ == '__main__':
    main()
