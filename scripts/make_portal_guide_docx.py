# -*- coding: utf-8 -*-
"""Generate Academor student portal Word guide (non-technical, AZ)."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

COVER = r"C:\Users\user\.cursor\projects\c-Users-user-Desktop-Academor\assets\portal_guide_cover.png"
ROLES = r"C:\Users\user\.cursor\projects\c-Users-user-Desktop-Academor\assets\portal_roles.png"
FLOW = r"C:\Users\user\.cursor\projects\c-Users-user-Desktop-Academor\assets\portal_flow.png"
OUT = r"C:\Users\user\Desktop\Academor\Academor_Sagird_Portali_Beledci.docx"

TEAL = RGBColor(0x0D, 0x6E, 0x6E)
NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GRAY = RGBColor(0x4A, 0x55, 0x68)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_TEAL = "E6F4F1"
LIGHT_BLUE = "E8F0FE"
LIGHT_GREEN = "E8F5E9"
LIGHT_AMBER = "FFF8E1"


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 18, True, TEAL)
    elif level == 2:
        set_run_font(run, 14, True, NAVY)
    else:
        set_run_font(run, 12, True, TEAL)
    return p


def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size, False, GRAY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, 11, False, GRAY)
    return p


def shade_cell(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=10, color=GRAY, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size, bold, color)


def add_callout(doc, title, body, fill=LIGHT_AMBER):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    set_run_font(r1, 11, True, NAVY)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_font(r2, 10, False, GRAY)
    doc.add_paragraph()


def add_caption(doc, text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(text)
    set_run_font(run, 9, False, GRAY)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Cover
    doc.add_picture(COVER, width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    set_run_font(title.add_run("Academor Şagird Portalı"), 26, True, TEAL)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run("Şagirdlər və müəllimlər üçün sadə bələdçi"), 14, False, NAVY)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta.add_run("Nə var? Necə işləyir? Nəyi bilmək lazımdır?"), 11, False, GRAY)

    doc.add_page_break()

    # 1
    add_heading_styled(doc, "1. Portala ümumi baxış", 1)
    add_body(
        doc,
        "Academor Şagird Portalı — dərslərin, tapşırıqların, davamiyyətin və nəticələrin "
        "bir yerdə toplandığı onlayn kabinetdir. Kompüterdən və ya telefondan daxil olmaq olar.",
    )
    add_body(
        doc,
        "Bu bələdçi texniki detallara girmədən izah edir: portalda nələr var, şagird nə edir, "
        "müəllim necə idarə edir və hər kəs nəyi bilməlidir.",
    )

    add_heading_styled(doc, "Kimlər istifadə edir?", 2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    roles = [
        (
            "Şagird",
            "Dərs cədvəlinə baxır, material oxuyur, ev tapşırığı göndərir, quiz və mock imtahan verir, nəticələrini izləyir.",
            LIGHT_TEAL,
        ),
        (
            "Müəllim",
            "Qrupları görür, davamiyyət qeyd edir, həftəlik qiymət yazır, material yükləyir, quiz və mock-ları yoxlayır.",
            LIGHT_BLUE,
        ),
        (
            "Valideyn",
            "Övladının cədvəlinə, dərslərinə, davamiyyətinə və nəticələrinə yalnız baxış rejimində baxır.",
            LIGHT_GREEN,
        ),
        (
            "Qeyd",
            "Hesablar və qruplar mərkəzdən yaradılır. Şagird və müəllimə istifadəçi adı və şifrə verilir.",
            LIGHT_AMBER,
        ),
    ]
    for i, (role, desc, fill) in enumerate(roles):
        set_cell_text(table.cell(i, 0), role, bold=True, size=11, color=TEAL, center=True)
        set_cell_text(table.cell(i, 1), desc, size=10)
        shade_cell(table.cell(i, 0), fill)

    doc.add_paragraph()
    doc.add_picture(ROLES, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Şəkil: Şagird və müəllimin əsas vəzifələri")

    # 2
    add_heading_styled(doc, "2. Necə daxil olmaq olar?", 1)
    add_body(doc, "Portala daxil olmaq üçün Academor-un sizə verdiyi istifadəçi adı və şifrə lazımdır.")
    add_bullet(doc, "Brauzerdə Academor saytını açın.")
    add_bullet(doc, "«Portal» və ya «Daxil ol» düyməsinə basın.")
    add_bullet(doc, "İstifadəçi adı və şifrənizi yazın.")
    add_bullet(doc, "Girişdən sonra öz kabinetiniz (Ana səhifə) açılır.")

    add_callout(
        doc,
        "İpucu",
        "Şifrəni unutmusunuzsa, Academor dəstəyi və ya müəlliminizlə əlaqə saxlayın. "
        "Hesabı başqa şəxslə paylaşmayın.",
    )

    doc.add_picture(FLOW, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Şəkil: Sadə yol — Giriş → Ana səhifə → Dərslər → Nəticə")

    # 3
    add_heading_styled(doc, "3. Şagird kabinetində nələr var?", 1)
    add_body(
        doc,
        "Soldakı menyu (telefondan aşağıdakı düymələr) vasitəsilə bölmələrə keçirsiniz. "
        "Hər bölmənin bir məqsədi var:",
    )

    features = [
        ("Ana səhifə", "Sizi salamlayır; bu günün dərsləri və ümumi göstəricilər (orta qiymət, davamiyyət, quiz) görünür."),
        ("Cədvəl", "Həftəlik dərs saatlarınızı göstərir — hansı gün, hansı saat, hansı qrup."),
        ("Davamiyyət", "İştirak etdiyiniz, gecikdiyiniz və ya iştirak etmədiyiniz günləri izləyirsiniz."),
        ("Dərslər", "Müəllimin yüklədiyi materiallar, fayllar və videolar. Ev tapşırığı varsa, buradan açıb cavab göndərə bilərsiniz."),
        ("Dərsliklər", "Kursunuza aid PDF və digər oxu materialları."),
        ("Nəticələr", "Quiz tarixçəsi, həftəlik qiymətlər (adətən 10 üzərindən) və mock imtahan nəticələri."),
        ("Quizlər", "Testlər — çoxseçimli, oxu, dinləmə, yazı və danışıq formatlarında."),
        ("Mock Test", "Tam IELTS və ya SAT imtahanı (yalnız kursunuzda varsa və müəllim açarsa)."),
        ("Bildirişlər", "Yeni qiymət, yoxlama, təklif və digər xəbərlər."),
        ("Profil", "Şəkil, qısa məlumat, telefon və əlaqə linkləri."),
    ]
    ft = doc.add_table(rows=1 + len(features), cols=2)
    ft.style = "Table Grid"
    set_cell_text(ft.cell(0, 0), "Bölmə", bold=True, size=11, color=WHITE, center=True)
    set_cell_text(ft.cell(0, 1), "Nə üçündür?", bold=True, size=11, color=WHITE)
    shade_cell(ft.cell(0, 0), "0D6E6E")
    shade_cell(ft.cell(0, 1), "0D6E6E")
    for i, (name, desc) in enumerate(features, start=1):
        set_cell_text(ft.cell(i, 0), name, bold=True, size=10, color=NAVY)
        set_cell_text(ft.cell(i, 1), desc, size=10)
        if i % 2 == 0:
            shade_cell(ft.cell(i, 0), LIGHT_TEAL)
            shade_cell(ft.cell(i, 1), LIGHT_TEAL)

    doc.add_paragraph()
    add_heading_styled(doc, "Şagird nəyi bilməlidir?", 2)
    add_bullet(doc, "Hər dərsdən əvvəl cədvələ baxın.")
    add_bullet(doc, "Dərslər bölməsində yeni material və ev tapşırığını yoxlayın; tapşırığı vaxtında göndərin.")
    add_bullet(doc, "Quizə başlamazdan əvvəl sakit mühit və stabil internet seçin — bəzi testlər vaxtlıdır.")
    add_bullet(doc, "Yazı və danışıq quizlərində cavab avtomatik qiymətlənmir; müəllim yoxlayana qədər gözləyin.")
    add_bullet(doc, "Mock imtahanı yalnız müəllim açdıqda görünür — bağlıdırsa, müəllimə müraciət edin.")
    add_bullet(doc, "Nəticələrə müntəzəm baxın; zəif tərəfləri müəllimlə müzakirə edin.")

    add_callout(
        doc,
        "Vacib",
        "Portalda mesajlaşma mərkəzi yoxdur. Suallarınız üçün müəllimə birbaşa "
        "(dərsdə və ya razılaşdırılmış əlaqə ilə) müraciət edin.",
        "FCE4EC",
    )

    # 4
    add_heading_styled(doc, "4. Quiz və mock imtahanlar", 1)
    add_body(doc, "Quizlər biliklərinizi yoxlamaq üçündür. Format fərqli ola bilər:")

    qt = doc.add_table(rows=6, cols=3)
    qt.style = "Table Grid"
    for j, h in enumerate(["Format", "Siz nə edirsiniz?", "Qiymət necə çıxır?"]):
        set_cell_text(qt.cell(0, j), h, bold=True, size=10, color=WHITE, center=True)
        shade_cell(qt.cell(0, j), "1A3A5C")

    quiz_rows = [
        ("Çoxseçimli", "Variantlardan düzgün cavabı seçirsiniz", "Avtomatik"),
        ("Oxu (Reading)", "Mətn oxuyub sualları cavablandırırsınız", "Avtomatik"),
        ("Dinləmə", "Audio dinləyib cavab yazırsınız", "Avtomatik / yoxlama"),
        ("Yazı (Writing)", "Esse və ya mətn yazırsınız", "Müəllim yoxlayır"),
        ("Danışıq", "Səs yazısı göndərirsiniz", "Müəllim yoxlayır"),
    ]
    for i, row in enumerate(quiz_rows, start=1):
        for j, val in enumerate(row):
            set_cell_text(qt.cell(i, j), val, bold=(j == 0), size=10, color=NAVY if j == 0 else GRAY)
        if i % 2 == 0:
            for j in range(3):
                shade_cell(qt.cell(i, j), LIGHT_BLUE)

    doc.add_paragraph()
    add_body(doc, "Mock Test — real imtahana bənzər tam sınaqdır:")
    add_bullet(doc, "IELTS: Dinləmə → Oxu → Yazı → Danışıq")
    add_bullet(doc, "SAT: Oxu & Yazı → Riyaziyyat")
    add_body(
        doc,
        "İmtahan bitəndən sonra nəticə Nəticələr bölməsində görünür. "
        "Yazı və danışıq hissələri müəllim yoxladıqdan sonra tamamlanır.",
    )

    # 5
    add_heading_styled(doc, "5. Müəllim portalı necə idarə olunur?", 1)
    add_body(
        doc,
        "Müəllim gündəlik işi portalda aparır. Qruplar, şagirdlər və əsas tənzimləmələr "
        "əvvəlcədən hazırlanır; müəllim isə dərs prosesini idarə edir.",
    )

    add_heading_styled(doc, "Müəllimin əsas işləri", 2)
    teacher_feats = [
        ("Qruplar", "Öz qruplarınızı və şagird siyahısını görürsünüz."),
        ("Şagird profili", "Hər şagird üçün bir mərkəz: davamiyyət, həftəlik qiymət, quiz tarixçəsi, mock nəticələri."),
        ("Davamiyyət", "Dərs sessiyasına görə iştirak / qeyb / gecikmə qeyd edirsiniz."),
        ("Həftəlik qiymət", "Həftədə bir dəfə (adətən 10 üzərindən) qiymət yazırsınız; saxlandıqdan sonra dəyişmək məhdud ola bilər."),
        ("Dərslər və dərsliklər", "Material, fayl və video yükləyirsiniz; şagirdlər dərhal görür."),
        ("Quiz icazəsi", "Şagird üçün quiz və ya mock-u aça və ya bağlaya bilərsiniz."),
        ("Yoxlama", "Yazı və danışıq cavablarını oxuyub qiymət və şərh verirsiniz."),
        ("Nəticələr", "Qrup üzrə ümumi mənzərəyə baxırsınız."),
    ]
    tt = doc.add_table(rows=1 + len(teacher_feats), cols=2)
    tt.style = "Table Grid"
    set_cell_text(tt.cell(0, 0), "Əməliyyat", bold=True, size=11, color=WHITE, center=True)
    set_cell_text(tt.cell(0, 1), "Qısa izah", bold=True, size=11, color=WHITE)
    shade_cell(tt.cell(0, 0), "2E7D32")
    shade_cell(tt.cell(0, 1), "2E7D32")
    for i, (name, desc) in enumerate(teacher_feats, start=1):
        set_cell_text(tt.cell(i, 0), name, bold=True, size=10, color=NAVY)
        set_cell_text(tt.cell(i, 1), desc, size=10)
        if i % 2 == 0:
            shade_cell(tt.cell(i, 0), LIGHT_GREEN)
            shade_cell(tt.cell(i, 1), LIGHT_GREEN)

    doc.add_paragraph()
    add_heading_styled(doc, "Müəllim nəyi bilməlidir?", 2)
    add_bullet(doc, "Dərsdən sonra davamiyyəti eyni gün qeyd edin — şagird və valideyn dərhal görür.")
    add_bullet(doc, "Həftəlik qiyməti vaxtında yazın; saxladıqdan sonra dəyişiklik çətin ola bilər.")
    add_bullet(doc, "Yeni materialı Dərslər / Dərsliklər bölməsinə yükləyin ki, şagird axtarmasın.")
    add_bullet(doc, "Quiz və mock-u yalnız hazır olan şagirdə açın.")
    add_bullet(doc, "Yazı və danışıq yoxlamalarını gecikdirməyin — şagird nəticəni gözləyir.")
    add_bullet(doc, "Şagird profilindən fərdi irəliləyişə baxın; zəif tərəfləri dərsdə müzakirə edin.")

    # 6
    add_heading_styled(doc, "6. Hamısı necə birləşir? (sadə sxem)", 1)
    add_body(doc, "Portalın iş prinsipi belədir:")
    for s in [
        "1. Hesab və qrup hazırlanır — şagird və müəllim giriş məlumatı alır.",
        "2. Müəllim material yükləyir, davamiyyət və qiymət yazır, quiz/mock açır.",
        "3. Şagird cədvələ baxır, dərs oxuyur, tapşırıq və quiz verir.",
        "4. Nəticələr hər kəsə görünür; valideyn də övladının irəliləyişini izləyə bilir.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(s), 11, False, GRAY)

    add_callout(
        doc,
        "Yadda saxlayın",
        "Portal öyrənməni asanlaşdırır, amma əvəz etmir. Uğur üçün: müntəzəm giriş, "
        "vaxtında tapşırıq, quizlərdə diqqət və müəllimlə əlaqə.",
        LIGHT_TEAL,
    )

    # 7
    add_heading_styled(doc, "7. Tez-tez verilən suallar", 1)
    faqs = [
        (
            "Niyə bəzi quizlər / Mock bağlıdır?",
            "Müəllim hələ açmayıb və ya kursunuzda bu imtahan yoxdur. Müəllimə müraciət edin.",
        ),
        (
            "Ev tapşırığını haradan göndərim?",
            "Dərslər bölməsindən müvafiq dərsi açın və tapşırıq formasından mətn və/və ya fayl göndərin.",
        ),
        (
            "Valideyn hər şeyi dəyişə bilərmi?",
            "Xeyr. Valideyn yalnız baxır; dəyişiklik etmir.",
        ),
        (
            "Telefonla işləyirmi?",
            "Bəli. Menyü telefonlarda aşağıda düymələr şəklində çıxır.",
        ),
        (
            "Şifrəmi kimə deyə bilərəm?",
            "Heç kimə. Yalnız Academor dəstəyi və ya rəsmi yolla bərpa edin.",
        ),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        set_run_font(p.add_run("S: " + q), 11, True, NAVY)
        p2 = doc.add_paragraph()
        set_run_font(p2.add_run("C: " + a), 11, False, GRAY)

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(end.add_run("— Academor Şagird Portalı bələdçisi —"), 10, True, TEAL)
    end2 = doc.add_paragraph()
    end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        end2.add_run("Suallarınız üçün müəlliminizə və ya Academor dəstəyinə müraciət edin."),
        9,
        False,
        GRAY,
    )

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    main()
